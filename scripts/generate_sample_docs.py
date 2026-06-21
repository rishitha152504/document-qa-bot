"""Generate sample source documents for the RAG pipeline."""

from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

DATA_DIR = Path(__file__).resolve().parent.parent / "data"


def write_business_pdf(path: Path) -> None:
    pages = [
        (
            "Acme Corporation — Annual Business Report FY2026\n\n"
            "Executive Summary\n"
            "Acme Corporation delivered strong financial performance in fiscal year 2026. "
            "Net revenue grew by 14% year-over-year to $2.4 billion. Operating margin "
            "improved to 18.5%, driven by cost optimization and product mix shifts toward "
            "higher-margin enterprise software subscriptions.\n\n"
            "Key Highlights\n"
            "- Total employees: 12,450 across 28 countries\n"
            "- Customer retention rate: 94%\n"
            "- New enterprise contracts signed: 312\n"
            "- R&D investment: $180 million (7.5% of revenue)"
        ),
        (
            "Financial Performance\n\n"
            "Revenue Breakdown by Segment:\n"
            "1. Cloud Services — $1.1B (46% of total revenue)\n"
            "2. Enterprise Software — $850M (35%)\n"
            "3. Professional Services — $450M (19%)\n\n"
            "Geographic Revenue:\n"
            "- North America: 52%\n"
            "- Europe: 28%\n"
            "- Asia-Pacific: 20%\n\n"
            "The Cloud Services segment showed the fastest growth at 22% YoY, "
            "while Professional Services grew 8% YoY."
        ),
        (
            "Strategic Initiatives & Outlook\n\n"
            "Acme launched Project Horizon in Q2 2026, an AI-powered analytics platform "
            "that integrates with existing enterprise workflows. Early adoption exceeded "
            "expectations with 1,200 beta customers in the first 90 days.\n\n"
            "FY2027 Guidance:\n"
            "- Expected revenue growth: 10-12%\n"
            "- Planned headcount increase: 800 employees\n"
            "- Capital expenditure: $95 million for data center expansion\n\n"
            "CEO Statement: 'Our focus remains on sustainable growth, customer success, "
            "and responsible AI deployment across all product lines.'"
        ),
    ]

    c = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter
    for page_text in pages:
        text_obj = c.beginText(50, height - 50)
        text_obj.setFont("Helvetica", 11)
        for line in page_text.split("\n"):
            text_obj.textLine(line)
        c.drawText(text_obj)
        c.showPage()
    c.save()


def write_science_pdf(path: Path) -> None:
    pages = [
        (
            "Advances in Neural Retrieval Systems — Research Paper\n\n"
            "Abstract\n"
            "This paper examines retrieval-augmented generation (RAG) architectures "
            "for domain-specific question answering. We compare dense vector retrieval "
            "against hybrid sparse-dense approaches across five benchmark datasets.\n\n"
            "Introduction\n"
            "Large language models excel at general reasoning but struggle with "
            "proprietary or time-sensitive knowledge. RAG addresses this by retrieving "
            "relevant document chunks at inference time and conditioning generation "
            "on retrieved context."
        ),
        (
            "Methodology\n\n"
            "We evaluated three chunking strategies:\n"
            "1. Fixed-size sliding window (500 tokens, 100 overlap)\n"
            "2. Recursive character splitting (1000 chars, 200 overlap)\n"
            "3. Semantic paragraph splitting\n\n"
            "Embedding models tested: text-embedding-004, all-MiniLM-L6-v2.\n"
            "Generation model: Gemini 2.0 Flash.\n\n"
            "Top-k retrieval values ranged from 3 to 10. Optimal performance was "
            "observed at k=5 with recursive character splitting."
        ),
        (
            "Results & Conclusion\n\n"
            "Recursive character splitting with k=5 achieved 87.3% answer accuracy "
            "on the proprietary document benchmark, outperforming fixed-size chunking "
            "by 4.2 percentage points.\n\n"
            "Key finding: chunk overlap of 150-200 characters significantly reduces "
            "boundary truncation errors. Strict grounding prompts reduced hallucination "
            "rates from 23% to 3% in controlled evaluations.\n\n"
            "Future work includes multi-hop retrieval and citation verification pipelines."
        ),
    ]

    c = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter
    for page_text in pages:
        text_obj = c.beginText(50, height - 50)
        text_obj.setFont("Helvetica", 11)
        for line in page_text.split("\n"):
            text_obj.textLine(line)
        c.drawText(text_obj)
        c.showPage()
    c.save()


def write_factsheet_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Product Factsheet — Acme Analytics Platform", level=1)
    doc.add_paragraph(
        "The Acme Analytics Platform is an enterprise-grade business intelligence "
        "solution designed for mid-to-large organizations. It supports real-time "
        "dashboards, automated reporting, and AI-assisted insights."
    )
    doc.add_heading("Core Features", level=2)
    doc.add_paragraph("Real-time data visualization with 50+ chart types")
    doc.add_paragraph("Natural language query interface powered by Gemini")
    doc.add_paragraph("Role-based access control with SSO integration")
    doc.add_paragraph("Export to PDF, Excel, and PowerPoint")
    doc.add_heading("Pricing", level=2)
    doc.add_paragraph("Starter Plan: $49/user/month (up to 25 users)")
    doc.add_paragraph("Business Plan: $89/user/month (unlimited users)")
    doc.add_paragraph("Enterprise Plan: Custom pricing with dedicated support")
    doc.add_heading("Support", level=2)
    doc.add_paragraph(
        "All plans include 24/7 email support. Business and Enterprise plans "
        "include phone support and a dedicated customer success manager."
    )
    doc.save(path)


def write_company_policy_txt(path: Path) -> None:
    path.write_text(
        "Acme Corporation — Remote Work Policy (Effective January 2026)\n\n"
        "1. Eligibility\n"
        "All full-time employees who have completed 90 days of employment are "
        "eligible for hybrid or fully remote work arrangements.\n\n"
        "2. Core Hours\n"
        "Employees must be available between 10:00 AM and 3:00 PM in their local "
        "time zone for meetings and collaboration.\n\n"
        "3. Equipment\n"
        "The company provides a laptop, monitor, and $500 home office stipend "
        "for eligible remote employees.\n\n"
        "4. Security\n"
        "All remote workers must use company VPN and enable multi-factor "
        "authentication on all corporate accounts.\n\n"
        "5. Review\n"
        "Remote work arrangements are reviewed annually by department managers.",
        encoding="utf-8",
    )


def write_technical_faq_txt(path: Path) -> None:
    path.write_text(
        "Acme Analytics Platform — Technical FAQ\n\n"
        "Q: What databases are supported?\n"
        "A: PostgreSQL, MySQL, SQL Server, Snowflake, BigQuery, and Redshift.\n\n"
        "Q: What is the maximum data refresh frequency?\n"
        "A: Real-time streaming for supported connectors; minimum batch interval is 5 minutes.\n\n"
        "Q: Does the platform support custom ML models?\n"
        "A: Yes. Enterprise plan customers can deploy custom models via the ML API endpoint.\n\n"
        "Q: What SLAs are offered?\n"
        "A: Business plan: 99.5% uptime. Enterprise plan: 99.9% uptime with financial credits.\n\n"
        "Q: How is data encrypted?\n"
        "A: AES-256 at rest, TLS 1.3 in transit. Keys managed via AWS KMS or Azure Key Vault.",
        encoding="utf-8",
    )


def main() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    write_business_pdf(DATA_DIR / "business_doc.pdf")
    write_science_pdf(DATA_DIR / "science_paper.pdf")
    write_factsheet_docx(DATA_DIR / "factsheet.docx")
    write_company_policy_txt(DATA_DIR / "company_policy.txt")
    write_technical_faq_txt(DATA_DIR / "technical_faq.txt")
    print(f"Sample documents created in {DATA_DIR}")


if __name__ == "__main__":
    main()
